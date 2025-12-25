"""根据 DOCX 讲稿和模板定义生成 JSON，可选调用 DeepSeek LLM。"""

from __future__ import annotations

import argparse
import json
import re
import secrets
import shutil
from datetime import datetime
from pathlib import Path
import os
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.oxml.ns import qn

from scripts.llm_client import (
    BaseLLM,
    DeepSeekLLM,
    GLMLLM,
    LocalLLM,
    QwenVLLM,
    TaichuLLM,
)
import base64
import mimetypes

MARKER_RE = re.compile(r"【PPT(\d+)】")
IMAGE_NAME_TEMPLATE = "doc_image_{idx}.{ext}"

# 调试标志：设置为 True 时打印 LLM 请求和响应
DEBUG_LLM = os.getenv("DEBUG_LLM", "false").lower() in ("true", "1", "yes")

# 章节标题正则表达式
# 一级标题：一、xxx / 第一章 xxx / 1 xxx
HEADING_L1_RE = re.compile(
    r"^(?:"
    r"[一二三四五六七八九十]+、|"  # 一、二、三、
    r"第[一二三四五六七八九十\d]+[章节部分]|"  # 第一章、第1节
    r"\d+\s+"  # 1 xxx（数字后有空格）
    r")(.+)",
    re.MULTILINE,
)
# 二级标题：（一）xxx / （1）xxx / 1.1 xxx
HEADING_L2_RE = re.compile(
    r"^(?:"
    r"[（\(][一二三四五六七八九十\d]+[）\)]|"  # （一）（1）
    r"\d+\.\d+\s*"  # 1.1 xxx
    r")(.+)",
    re.MULTILINE,
)
# 三级标题：1. xxx / ① xxx / a) xxx
HEADING_L3_RE = re.compile(
    r"^(?:"
    r"\d+\.\s+|"  # 1. xxx
    r"[①②③④⑤⑥⑦⑧⑨⑩]|"  # ① xxx
    r"[a-z][）\)]\s*"  # a) xxx
    r")(.+)",
    re.MULTILINE,
)


def _extract_heading_level(text: str) -> Tuple[Optional[int], Optional[str]]:
    """
    从文本中提取章节标题级别和标题内容。

    Returns:
        (level, title): level 为 1/2/3 表示一/二/三级标题，None 表示不是标题
    """
    text = text.strip()
    if not text:
        return None, None

    # 检查一级标题
    match = HEADING_L1_RE.match(text)
    if match:
        # 返回完整的标题文本（包括序号）
        return 1, text

    # 检查二级标题
    match = HEADING_L2_RE.match(text)
    if match:
        return 2, text

    # 检查三级标题
    match = HEADING_L3_RE.match(text)
    if match:
        return 3, text

    return None, None


class SectionContext:
    """章节上下文，跟踪当前所处的章节层级"""

    def __init__(self):
        self.level1: Optional[str] = None  # 一级标题（章）
        self.level2: Optional[str] = None  # 二级标题（节/知识点）
        self.level3: Optional[str] = None  # 三级标题（小节）

    def update(self, text: str) -> None:
        """根据文本内容更新章节上下文，只识别明确的章节标题格式"""
        level, title = _extract_heading_level(text)
        if level == 1:
            self.level1 = title
            self.level2 = None  # 重置下级标题
            self.level3 = None
        elif level == 2:
            self.level2 = title
            self.level3 = None  # 重置下级标题
        elif level == 3:
            self.level3 = title
        # 不再使用兜底逻辑，避免把无关内容当作章节标题

    def to_dict(self) -> Dict[str, Optional[str]]:
        """返回当前章节上下文的字典表示"""
        return {
            "chapter": self.level1,  # 章节名称
            "section": self.level2,  # 知识点名称
            "subsection": self.level3,  # 小节名称
        }

    def __str__(self) -> str:
        parts = []
        if self.level1:
            parts.append(f"章节：{self.level1}")
        if self.level2:
            parts.append(f"知识点：{self.level2}")
        if self.level3:
            parts.append(f"小节：{self.level3}")
        return " > ".join(parts) if parts else "（无章节信息）"


def _create_run_dir(base_dir: Path = Path("temp")) -> Path:
    """创建带时间戳前缀的运行目录，方便前端一次处理对应到单独目录。"""
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    suffix = secrets.token_hex(2)
    run_dir = base_dir / f"script-{timestamp}-{suffix}"
    run_dir.mkdir(parents=True, exist_ok=True)
    return run_dir


def parse_docx_blocks(doc_path: str, image_dir: Path) -> Tuple[List[Dict], bool, Dict]:
    """读取 DOCX 并按照 PPT 标记拆分内容，同时保存提取的图片与课程元信息。"""
    doc = Document(doc_path)
    image_dir.mkdir(parents=True, exist_ok=True)
    slides: List[Dict] = []
    current: Optional[Dict] = None
    buffer: List[str] = []
    has_marker = False
    image_counter = 1
    metadata: Dict[str, str] = {}

    def flush():
        nonlocal current, buffer
        if current is None and buffer:
            slides.append(
                {"template_hint": None, "text": "\n".join(buffer).strip(), "images": []}
            )
        elif current:
            current["text"] = "\n".join(buffer).strip()
            slides.append(current)
        buffer = []
        current = None

    def ensure_block():
        nonlocal current
        if current is None:
            current = {"template_hint": None, "text": "", "images": []}

    def attach_images(paths: List[str]):
        if not paths:
            return
        ensure_block()
        current.setdefault("images", []).extend(paths)

    def extract_images(paragraph) -> List[str]:
        nonlocal image_counter
        images = []
        for element in paragraph._p.iter():
            if element.tag not in {qn("a:blip"), qn("pic:blip")}:
                continue
            r_id = element.get(qn("r:embed"))
            if not r_id:
                continue
            part = paragraph.part.related_parts.get(r_id)
            if not part:
                continue
            ext = part.filename.split(".")[-1].lower() or "png"
            filename = image_dir / IMAGE_NAME_TEMPLATE.format(
                idx=image_counter, ext=ext
            )
            with open(filename, "wb") as f:
                f.write(part.blob)
            images.append(str(filename))
            image_counter += 1
        return images

    for para in doc.paragraphs:
        text = para.text
        stripped = text.strip()
        meta_match = re.match(r"^(课程名称|学院名称|主讲教师)[：:]\s*(.+)$", stripped)
        if meta_match:
            key = meta_match.group(1)
            value = meta_match.group(2).strip()
            if key == "课程名称":
                metadata["course"] = value
            elif key == "学院名称":
                metadata["college"] = value
            elif key == "主讲教师":
                metadata["lecturer"] = value
            continue
        images = extract_images(para)
        if images:
            attach_images(images)
            # 不在 buffer 中添加图片标记，图片路径只保存在 block["images"] 中
            # 图片标记会在 _preprocess_and_fill 中统一添加

        matches = list(MARKER_RE.finditer(text))
        if matches:
            idx = 0
            for match in matches:
                prefix = text[idx : match.start()].strip()
                if prefix:
                    buffer.append(prefix)
                flush()
                has_marker = True
                current = {
                    "template_hint": int(match.group(1)),
                    "text": "",
                    "images": [],
                }
                idx = match.end()
            remainder = text[idx:].strip()
            if remainder:
                buffer.append(remainder)
            continue

        if stripped:
            buffer.append(stripped)

    flush()
    if not slides:
        full_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        slides.append({"template_hint": None, "text": full_text, "images": []})
    return slides, has_marker, metadata


def load_template_defs(
    template_json: str, template_list: Optional[str] = None
) -> Dict[int, Dict]:
    """
    加载模板定义。

    Args:
        template_json: template.json 文件路径
        template_list: （已废弃）template.txt 文件路径，用于过滤允许的模板页码
                      如果文件不存在或内容为空，则不进行过滤
    """
    data = json.loads(Path(template_json).read_text(encoding="utf-8"))
    allowed = None
    # template.txt 已废弃，但为了兼容旧版本，仍然支持
    if template_list and Path(template_list).exists():
        content = Path(template_list).read_text(encoding="utf-8").strip()
        if content:  # 只有内容非空时才过滤
            allowed = {
                int(item.strip())
                for item in content.replace(",", " ").split()
                if item.strip().isdigit()
            }
    templates = {}
    manifest = {
        item["template_page_num"]: item
        for item in data.get("manifest", [])
        if "template_page_num" in item
    }
    for page in data.get("ppt_pages", []):
        num = page.get("template_page_num")
        if not isinstance(num, int):
            continue
        if allowed and num not in allowed:
            continue
        schema = page.get("content", {})
        fields = _collect_fields(schema)
        # 优先使用顶层 page_note，其次使用 meta.notes
        meta = page.get("meta", {}) or manifest.get(num, {})
        page_note = page.get("page_note") or meta.get("notes", "")
        templates[num] = {
            "page_type": page.get("page_type", f"模板第{num}页"),
            "schema": schema,
            "meta": meta,
            "page_note": page_note,  # 统一使用 page_note
            "text_fields": [field for field in fields if not field["is_image"]],
            "image_fields": [field for field in fields if field["is_image"]],
        }
    if not templates:
        raise ValueError("模板列表为空，无法匹配。")
    return templates


def _collect_fields(schema, prefix=None):
    results = []
    prefix = prefix or ()
    if isinstance(schema, dict):
        if "type" in schema and "value" in schema:
            field_type = schema.get("type") or "text"
            results.append(
                {
                    "path": prefix,
                    "is_image": field_type.lower() == "image",
                    "hint": schema.get("hint") or "",
                    "max_chars": schema.get("max_chars"),
                    "required": bool(schema.get("required", False)),
                }
            )
        else:
            for key, value in schema.items():
                results.extend(_collect_fields(value, prefix + (key,)))
    else:
        # 兼容旧格式：字符串或其他类型视为文本叶子
        results.append(
            {
                "path": prefix,
                "is_image": any(
                    "图片" in seg or "image" in seg.lower() for seg in prefix
                ),
                "hint": "",
                "max_chars": None,
                "required": False,
            }
        )
    return results


def _clone_schema(schema: Dict) -> Dict:
    return json.loads(json.dumps(schema, ensure_ascii=False))


def _assign_in_schema(schema: Dict, path: List[str], value: str):
    node = schema
    for key in path[:-1]:
        node = node.setdefault(key, {})
    leaf = node.get(path[-1])
    if isinstance(leaf, dict) and "type" in leaf:
        leaf["value"] = value
    else:
        node[path[-1]] = value


def _is_multimodal_llm(llm: Optional[BaseLLM]) -> bool:
    """检查是否为多模态模型"""
    return isinstance(llm, (TaichuLLM, GLMLLM))


def _simple_fill(template_info: Dict, raw_text: str, images: List[str]) -> Dict:
    result = _clone_schema(template_info["schema"])
    text_fields = template_info["text_fields"]
    image_fields = template_info["image_fields"]
    if text_fields:
        _assign_in_schema(result, list(text_fields[0]["path"]), raw_text)
        for field in text_fields[1:]:
            _assign_in_schema(result, list(field["path"]), "")
    for idx, field in enumerate(image_fields):
        value = images[idx] if idx < len(images) else ""
        _assign_in_schema(result, list(field["path"]), value)
    return result


def _build_prompt(
    template_info: Dict,
    raw_text: str,
    images: List[str],
    is_multimodal: bool = False,
    user_prompt: Optional[str] = None,
    metadata: Optional[Dict] = None,
    section_context: Optional[Dict[str, Optional[str]]] = None,
) -> str:
    """
    构建 LLM prompt，核心原则：
    - hint 是填充该字段的最重要依据
    - required 和 max_chars 是必须遵守的约束
    - page_note 是本页的特殊说明，必须优先遵守
    """

    def describe_fields(fields, is_image: bool = False):
        """详细描述每个字段，突出 hint 的指导作用"""
        if not fields:
            return "（无）"
        lines = []
        for idx, field in enumerate(fields, 1):
            name = "/".join(field["path"])
            hint = field.get("hint") or ""
            max_chars = field.get("max_chars")
            required = field.get("required", False)

            # 构建字段描述，hint 作为核心指导
            parts = [f"  【字段{idx}】{name}"]
            if hint:
                parts.append(f"    → 填写要求：{hint}")
            if required:
                parts.append(f"    → ⚠️ 必填")
            else:
                parts.append(f"    → 可选（无匹配内容时留空）")
            if max_chars and not is_image:
                parts.append(f"    → 字数限制：≤{max_chars}字")

            lines.append("\n".join(parts))
        return "\n\n".join(lines)

    text_desc = describe_fields(template_info["text_fields"], is_image=False)
    image_desc = describe_fields(template_info["image_fields"], is_image=True)

    # 对于多模态模型，不需要列出图片路径
    if is_multimodal:
        image_section = f"已附加 {len(images)} 张图片"
    else:
        image_section = "无" if not images else "\n".join(images)

    # page_note 是本页最重要的特殊说明
    page_note = template_info.get("page_note") or ""
    meta = template_info.get("meta") or {}
    if not page_note:
        page_note = meta.get("notes", "")

    # 构建 page_note 部分，如果有内容则高亮显示
    page_note_section = ""
    if page_note:
        page_note_section = f"""
════════════════════════════════════════
📋 【本页特殊说明 - 必须优先遵守】
{page_note}
════════════════════════════════════════
"""

    # 多模态图片说明
    multimodal_instruction = ""
    if is_multimodal and images:
        multimodal_instruction = f"""
🖼️ 图片处理说明：
已附加 {len(images)} 张图片。图片与文本的关系：
- 每张图片紧跟在相关文本段落的下方，是对上方文本的补充说明
- 讲稿中的 `[图片资源: ...]` 标记指示图片在原文中的位置
- 请根据图片字段的 hint 要求，将图片填入对应字段
- 如果 page_note 中有图片选择规则（如"选择1,3"），必须严格遵守
"""

    # 预设元信息
    metadata_section = ""
    if metadata and any(metadata.values()):
        meta_items = []
        if metadata.get("course"):
            meta_items.append(f"课程名称：{metadata['course']}（填入 hint 中要求'课程'、'项目'的字段）")
        if metadata.get("college"):
            meta_items.append(f"学院名称：{metadata['college']}（填入 hint 中要求'学院'、'单位'的字段）")
        if metadata.get("lecturer"):
            meta_items.append(f"主讲人：{metadata['lecturer']}（填入 hint 中要求'主讲'、'讲师'、'姓名'的字段）")
        if meta_items:
            metadata_section = f"""
📌 预设信息（严格按照括号内的指示填入对应字段）：
{chr(10).join('- ' + item for item in meta_items)}
⚠️ 注意：主讲人姓名只能填入 hint 明确要求'主讲人'或'姓名'的字段，绝对不能填入'一级标题'、'章节标题'等标题类字段！
"""

    # 章节上下文
    context_section = ""
    if section_context and any(section_context.values()):
        ctx_items = []
        if section_context.get("chapter"):
            ctx_items.append(f"章节：{section_context['chapter']}（填入 hint 中要求'章节'、'一级标题'的字段）")
        if section_context.get("section"):
            ctx_items.append(f"知识点：{section_context['section']}（填入 hint 中要求'知识点'、'二级标题'的字段）")
        if section_context.get("subsection"):
            ctx_items.append(f"小节：{section_context['subsection']}（填入 hint 中要求'小节'、'三级标题'的字段）")
        if ctx_items:
            context_section = f"""
📚 当前章节上下文（按照括号内的指示填入对应字段）：
{chr(10).join('- ' + item for item in ctx_items)}
"""

    prompt = f"""请根据讲稿内容填充模板《{template_info["page_type"]}》的各个字段。
{page_note_section}
═══════════════════════════════════════
📝 【字段定义 - 严格按照 hint 要求填写】

每个字段的 hint 是填写该字段的最重要依据，必须严格遵守。

【文本字段】
{text_desc}

【图片字段】
{image_desc}

可用图片：{image_section}
════════════════════════════════════════
{metadata_section}{context_section}{multimodal_instruction}
📌 填写原则：
1. 每个字段严格按照其 hint 要求填写，hint 是唯一的填写依据
2. 必填字段不得留空，可选字段无匹配内容时留空
3. 严格遵守字数限制，超出限制的内容需精简
4. 从讲稿中提炼要点，不要照搬原文
5. 如果 page_note 中有特殊规则（如图片选择规则），必须严格遵守

输出格式：
{{
  "texts": ["字段1内容", "字段2内容", ...],
  "images": ["图片路径1", "图片路径2", ...]
}}

数组顺序必须与上述字段定义顺序一致。

讲稿内容：
{raw_text}
"""
    if user_prompt:
        prompt += f"\n\n用户额外要求：\n{user_prompt}"

    return prompt


def _encode_image(image_path: str) -> Optional[str]:
    """Read image file and return base64 string.

    Returns:
        Base64 encoded string, or None if encoding fails.
    """
    try:
        with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode("utf-8")
    except FileNotFoundError:
        print(f"⚠️ 警告：图片文件不存在：{image_path}")
        return None
    except Exception as e:
        print(f"⚠️ 警告：读取图片失败 {image_path}：{e}")
        return None


def _build_multimodal_messages(
    template_info: Dict,
    raw_text: str,
    images: List[str],
    user_prompt: Optional[str] = None,
    metadata: Optional[Dict] = None,
    section_context: Optional[Dict[str, Optional[str]]] = None,
) -> List[Dict]:
    """构建多模态消息，用于 Taichu-VL 或 GLMV。"""
    prompt_text = _build_prompt(
        template_info,
        raw_text,
        images,
        is_multimodal=True,
        user_prompt=user_prompt,
        metadata=metadata,
        section_context=section_context,
    )

    content: List[Dict[str, Any]] = [{"type": "text", "text": prompt_text}]

    # 添加图片
    for img_path in images:
        if not os.path.exists(img_path):
            print(f"⚠️ 警告：跳过不存在的图片：{img_path}")
            continue

        # Taichu-VL 使用 OpenAI 兼容格式，支持 data URL (base64)
        # 参考：https://docs.wair.ac.cn/intelligent/maas/visioIntro.html
        mime_type, _ = mimetypes.guess_type(img_path)
        if not mime_type:
            mime_type = "image/jpeg"

        base64_str = _encode_image(img_path)
        if not base64_str:  # 编码失败，跳过此图片
            continue

        data_url = f"data:{mime_type};base64,{base64_str}"

        content.append({"type": "image_url", "image_url": {"url": data_url}})

    return [{"role": "user", "content": content}]


def _lookup_field_value(field, payload, fallback_list, idx):
    path = list(field["path"])
    key = "/".join(path)

    def normalize(text):
        return "".join(ch for ch in text.lower() if ch not in {" ", "_"})

    if isinstance(payload, dict):
        norm_targets = {normalize(key), normalize(path[-1])}
        for candidate_key, candidate_val in payload.items():
            norm = normalize(candidate_key)
            if norm in norm_targets or any(
                target and target in norm for target in norm_targets
            ):
                return candidate_val
        values = list(payload.values())
        if idx < len(values):
            return values[idx]
        return values[-1] if values else ""

    if isinstance(fallback_list, list):
        return fallback_list[idx] if idx < len(fallback_list) else ""
    return ""


def llm_fill_slide(
    llm: BaseLLM,
    template_info: Dict,
    raw_text: str,
    images: List[str],
    user_prompt: Optional[str] = None,
    use_multimodal: bool = True,
    metadata: Optional[Dict] = None,
    section_context: Optional[Dict[str, Optional[str]]] = None,
) -> Dict:
    """
    使用 LLM 填充单页幻灯片内容。

    Args:
        llm: LLM 实例
        template_info: 模板信息
        raw_text: 原始文本
        images: 图片路径列表
        user_prompt: 用户自定义 prompt
        use_multimodal: 是否使用多模态消息（默认 True）
                       当讲稿有 PPT 标记时，图片位置已确定，可设为 False
        metadata: 元数据（课程名称、学院名称、讲师名称）
        section_context: 章节上下文（chapter/section/subsection）
    """
    if not llm:
        return _simple_fill(template_info, raw_text, images)

    # 只有在允许使用多模态且模型支持多模态且有图片时，才使用多模态消息
    if use_multimodal and _is_multimodal_llm(llm) and images:
        messages = _build_multimodal_messages(
            template_info, raw_text, images, user_prompt, metadata, section_context
        )
    else:
        prompt = _build_prompt(
            template_info,
            raw_text,
            images,
            user_prompt=user_prompt,
            metadata=metadata,
            section_context=section_context,
        )
        messages = [{"role": "user", "content": prompt}]

    if DEBUG_LLM:
        print(f"\n{'=' * 60}")
        print(f"🔍 [DEBUG] LLM 请求 (llm_fill_slide)")
        print(f"{'=' * 60}")
        # 检查实际发送的消息类型
        is_multimodal_message = messages and isinstance(
            messages[0].get("content"), list
        )
        if is_multimodal_message:
            print(f"📝 多模态消息 (文本 + {len(images)} 张图片)")
            # 只打印文本部分，图片太长不打印
            for msg in messages:
                if isinstance(msg.get("content"), list):
                    for item in msg["content"]:
                        if item.get("type") == "text":
                            print(f"文本内容:\n{item['text'][:500]}...")
        else:
            print(f"📝 文本消息:\n{messages[0]['content'][:500]}...")
        print(f"{'=' * 60}\n")

    response = llm.generate(messages, temperature=0.2)

    if DEBUG_LLM:
        print(f"\n{'=' * 60}")
        print(f"📥 [DEBUG] LLM 响应 (llm_fill_slide)")
        print(f"{'=' * 60}")
        print(f"{response[:500]}...")
        print(f"{'=' * 60}\n")
    try:
        data = _ensure_json_object(response)
        texts = data.get("texts", [])
        imgs = data.get("images", [])
    except Exception:
        texts, imgs = [raw_text], images

    result = _clone_schema(template_info["schema"])
    text_fields = template_info["text_fields"]
    image_fields = template_info["image_fields"]

    for idx, field in enumerate(text_fields):
        value = _lookup_field_value(
            field, texts, texts if isinstance(texts, list) else None, idx
        )
        _assign_in_schema(result, list(field["path"]), value)

    for idx, field in enumerate(image_fields):
        value = _lookup_field_value(
            field, imgs, imgs if isinstance(imgs, list) else None, idx
        )
        if not value and isinstance(images, list) and idx < len(images):
            value = images[idx]
        _assign_in_schema(result, list(field["path"]), value)

    return result


def llm_plan_slides(
    llm: BaseLLM,
    doc_text: str,
    templates: Dict[int, Dict],
    images: List[str],
    user_prompt: Optional[str] = None,
) -> List[Dict]:
    template_desc = "\n".join(
        f"- 模板 {info['page_type']} (编号 {num}): 文本{len(info['text_fields'])}项, 图片{len(info['image_fields'])}项"
        for num, info in templates.items()
    )

    # 对于多模态模型，不需要列出图片路径（图片已通过 base64 附加）
    if _is_multimodal_llm(llm) and images:
        image_section = f"已附加 {len(images)} 张图片供你参考"
    else:
        image_section = "无" if not images else "\n".join(images)

    multimodal_instruction = ""
    if _is_multimodal_llm(llm) and images:
        multimodal_instruction = f"""
🖼️ 多模态图片理解（重要）：
我已附带了 {len(images)} 张图片，这些图片是讲稿的重要组成部分，请务必认真处理。

图片与文本的关系：
- 每张图片都紧跟在相关文本段落的下方（图片在文本下方）
- 图片是对上方文本的补充说明、示例或可视化
- 讲稿文本中的 `[图片资源: ...]` 标记仅用于指示图片在原文中的位置

你的任务：
1. 仔细查看每张图片的内容，理解图片传达的信息
2. 分析图片与上下文文本的关系，确定图片所属的主题
3. 将图片放入合适的模板的图片字段（通常与相关文本在同一页 PPT）
4. 在输出的 JSON 对象中，"images" 数组应包含你选择使用的图片完整路径
5. 根据图片内容优化文本描述，使其更准确、更生动

⚠️ 重要：不要忽略图片！图片是讲稿的核心内容之一，必须合理使用。
"""

    prompt = f"""
请将以下讲稿拆分成若干张 PPT，每张幻灯片选择一个模板，并输出 JSON 数组，每个元素包含：
- template_page_num: 模板编号
- page_type: 模板名称
- texts: 按模板文本字段顺序给出的内容数组
- images: 按模板图片字段顺序给出的内容数组

模板信息：
{template_desc}

可用图片路径：
{image_section}

{multimodal_instruction}

输出格式示例：
[
  {{
    "template_page_num": 4,
    "page_type": "目录页",
    "texts": ["目录标题", "条目1", "条目1说明", "..."],
    "images": [""]
  }},
  ...
]


📌 核心原则（幻灯片 vs 讲稿）：
讲稿是演讲者手里的稿子，是他演讲时要说的完整内容。
幻灯片是投影给观众看的，应该是讲稿的**精炼要点**，而非照搬全文。
你需要把讲稿内容**提炼、概括、分点**后放到幻灯片上。

✅ 正确做法：
- 提取核心观点，用简洁的短语或短句表达
- 使用要点列表（如"1. xxx  2. xxx"或"• xxx"）
- 删除口语化表达、过渡语、详细解释
- 保留关键数据、专有名词、核心结论

❌ 错误做法：
- 把讲稿的长段落直接复制到幻灯片
- 保留"接下来我们来看""正如前面所说"等口语
- 内容过于详细，像在读文章

⚠️ 严格要求（必须遵守）：
1. 所有标记为"required"的字段必须填写，绝对不得留空。
2. 每个文本字段都有字数上限（max_chars），你生成的内容绝对不能超过这个限制。
3. 内容必须精炼！把讲稿的完整表述压缩为幻灯片要点，只保留核心信息。
4. 违反字数限制的输出将被视为无效，必须重新生成。
5. 务必记住并重复使用讲稿中的主讲人姓名、课程/讲座/项目名称等关键专有名词，确保在所有幻灯片中需要填写专有名词的位置保持一致，不要随意改写或另造新名称。
6. 同一页内的不同字段要符合各自语义，严禁把同一句话复制到多个字段，或用课程名/章节名去填正文/小节标题；标题/正文/要点请各填最贴合的内容。

讲稿全文：
{doc_text}
"""
    # Append user prompt if provided
    if user_prompt:
        prompt += f"\n\n用户额外要求：\n{user_prompt}"

    if _is_multimodal_llm(llm) and images:
        content: List[Dict[str, Any]] = [{"type": "text", "text": prompt}]
        for img_path in images:
            if not os.path.exists(img_path):
                print(f"⚠️ 警告：跳过不存在的图片：{img_path}")
                continue
            mime_type, _ = mimetypes.guess_type(img_path)
            if not mime_type:
                mime_type = "image/jpeg"
            base64_str = _encode_image(img_path)
            if not base64_str:  # 编码失败，跳过此图片
                continue
            data_url = f"data:{mime_type};base64,{base64_str}"
            content.append({"type": "image_url", "image_url": {"url": data_url}})
        messages = [{"role": "user", "content": content}]
    else:
        messages = [{"role": "user", "content": prompt}]

    if DEBUG_LLM:
        print(f"\n{'=' * 60}")
        print("🔍 [DEBUG] LLM 请求 (llm_plan_slides)")
        print(f"{'=' * 60}")
        if _is_multimodal_llm(llm) and images:
            print(f"📝 多模态消息 (文本 + {len(images)} 张图片)")
            # 只打印文本部分
            for msg in messages:
                if isinstance(msg.get("content"), list):
                    for item in msg["content"]:
                        if isinstance(item, dict) and item.get("type") == "text":
                            print(f"文本内容:\n{item['text'][:500]}...")
        else:
            print(f"📝 文本消息:\n{messages[0]['content'][:500]}...")
        print(f"{'=' * 60}\n")

    response = llm.generate(messages, temperature=0.3)

    if DEBUG_LLM:
        print(f"\n{'=' * 60}")
        print("📥 [DEBUG] LLM 响应 (llm_plan_slides)")
        print(f"{'=' * 60}")
        print(f"{response[:500]}...")
        print(f"{'=' * 60}\n")

    try:
        plan = _ensure_json_array(response)
        return plan
    except Exception:
        raise ValueError("模型输出无法解析为 JSON 数组，请检查提示或重试。")


def llm_preprocess_script(
    llm: BaseLLM,
    doc_text: str,
    templates: Dict[int, Dict],
    images: List[str],
    user_prompt: Optional[str] = None,
    has_marker: bool = False,
) -> str:
    """
    使用 LLM 预处理讲稿。

    Args:
        llm: LLM 实例
        doc_text: 原始讲稿文本
        templates: 模板定义字典
        images: 图片路径列表
        user_prompt: 用户自定义提示
        has_marker: 讲稿是否已有【PPT】标记

    Returns:
        带有【PPT1】【PPT2】等标记的 Markdown 格式讲稿
    """
    if not llm:
        raise ValueError("预处理讲稿需要启用 LLM。")

    # 构建模板描述，突出 page_note 和字段 hint
    template_desc_lines = []
    for num, info in templates.items():
        text_count = len(info["text_fields"])
        image_count = len(info["image_fields"])
        page_type = info["page_type"]
        has_required_image = any(f.get("required") for f in info["image_fields"])

        # 获取 page_note
        page_note = info.get("page_note") or ""
        meta = info.get("meta") or {}
        if not page_note:
            page_note = meta.get("notes", "")

        # 获取文本字段的详细信息（包含 hint）
        text_fields_desc = []
        for field in info["text_fields"]:
            path = field.get("path", ())
            name = "/".join(path) if path else "未命名"
            hint = field.get("hint", "")
            max_chars = field.get("max_chars")
            required = field.get("required", False)
            req_mark = "必填" if required else "可选"
            char_limit = f"≤{max_chars}字" if max_chars else ""
            desc_parts = [name]
            if hint:
                desc_parts.append(f"({hint})")
            if char_limit:
                desc_parts.append(f"[{char_limit}]")
            desc_parts.append(f"[{req_mark}]")
            text_fields_desc.append(f"    - {' '.join(desc_parts)}")

        image_fields_desc = []
        for field in info["image_fields"]:
            path = field.get("path", ())
            name = "/".join(path) if path else "图片"
            hint = field.get("hint", "")
            required = "必填" if field.get("required") else "可选"
            desc_parts = [name]
            if hint:
                desc_parts.append(f"({hint})")
            desc_parts.append(f"[{required}]")
            image_fields_desc.append(f"    - {' '.join(desc_parts)}")

        # 构建模板描述
        desc = f"【PPT{num}】{page_type}"
        if has_required_image:
            desc += " ⚠️需要图片"
        
        # 如果有 page_note，高亮显示
        if page_note:
            desc += f"\n  📋 特殊说明：{page_note}"
        
        desc += f"\n  文本字段({text_count}个):"
        if text_fields_desc:
            desc += "\n" + "\n".join(text_fields_desc)
        else:
            desc += "\n    （无）"
        
        if image_count > 0:
            desc += f"\n  图片字段({image_count}个):"
            desc += "\n" + "\n".join(image_fields_desc)
        
        template_desc_lines.append(desc)

    template_desc = "\n\n".join(template_desc_lines)

    # 根据是否已有标记，构建不同的 prompt
    if has_marker:
        # 已有标记：保持分页，只优化文本
        prompt = f"""你是一位专业的演讲稿编辑。讲稿已经分好页了，请**保持原有分页结构**，只优化文本表达。

## 任务说明

1. **保持分页不变**：讲稿中的【PPT编号】标记表示分页，必须原样保留，不能修改、删除或重新分配
2. **优化文本表达**：将口语化表达改为正式书面语，但不改变语义
3. **保持内容完整**：不要精简或删减内容，保持原文语义完整

⚠️ **重要**：分页已由用户指定，你只能优化文本，不能改变分页！

## 可用模板（仅供参考）

{template_desc}

## 注意事项

1. 【PPT编号】标记必须原样保留，不能修改
2. 不要精简内容，保持讲稿原文的完整性
3. 保留所有图片引用 `[图片资源: ...]`，位置不变
4. 不要输出任何解释说明，只输出优化后的讲稿

## 原始讲稿

{doc_text}
"""
    else:
        # 无标记：自动分页
        # 图片信息和模板限制
        if images:
            image_info = f"讲稿中包含 {len(images)} 张图片，请在适当位置保留图片引用。"
        else:
            image_info = "⚠️ 讲稿中【没有图片】！"
            # 找出需要图片的模板
            templates_need_image = [
                num
                for num, info in templates.items()
                if any(f.get("required") for f in info["image_fields"])
            ]
            if templates_need_image:
                image_info += (
                    f"\n【禁止】使用以下需要图片的模板：{templates_need_image}"
                )

        prompt = f"""你是一位专业的演讲稿编辑。请将以下原始讲稿**分页**，为每个部分选择合适的 PPT 模板。

## 任务说明

1. **分析讲稿结构**：理解讲稿的主题、逻辑和内容层次
2. **选择合适模板**：根据内容为每个部分选择最合适的 PPT 模板
3. **添加页码标记**：在每个部分开头用【PPT编号】标记该部分使用的模板
4. **保持内容完整**：不要精简或删减内容，保持原文语义完整，只需将口语化表达改为正式书面语

⚠️ **重要**：这一步只负责**分页和选择模板**，不要精简内容！

## 可用模板

{template_desc}

## 章节/知识点抽取指引
- 如果原文未显式写出章节/知识点，请根据语义推断出“章节/知识点”层级，并在对应页开头写清楚（如“【PPT2】\n常用传感器及其工作原理简介\n（一）光电传感器”）。
- 同一层级的标题保持用词一致，后续页面复用相同的章节/知识点名称，不要随意改写。
- 仅在有合理子标题时写三级标题；无合适小节时可以不写三级标题，避免用课程名/正文句子填充。

## 图片处理规则

{image_info}

**图片与模板匹配规则**：
- 图片标记 `[图片资源: xxx]` 必须与其**紧邻的上方文本**分到同一页
- 如果某页内容包含图片，应选择带图片字段的模板
- 如果某页内容**没有图片**，**禁止**选择标有"⚠️需要图片"的模板
- 图片是对上方文本的说明或示例，不能与文本分离

## 输出格式要求

输出为 Markdown 格式，每个 PPT 页面以【PPT编号】开头，例如：

```
【PPT2】
# 课程介绍

本课程将带您了解人工智能的基础知识，涵盖机器学习、深度学习等核心领域。

【PPT5】
# 机器学习基础

机器学习是人工智能的核心技术，它让计算机能够从数据中学习规律。

[图片资源: doc_image_1.png]
```

## 注意事项

1. 每个【PPT编号】标记必须独占一行
2. 编号必须是上面模板列表中存在的编号
3. **没有图片的页面禁止使用需要图片的模板**
4. 不要精简内容，保持讲稿原文的完整性
5. 图片引用 `[图片资源: ...]` 必须与上方文本保持在同一页
6. 不要输出任何解释说明，只输出分页后的讲稿
7. 章节/知识点/小节标题用词前后一致，无合适小节可不写三级标题，禁止用课程名或正文句子充当标题。

## 原始讲稿

{doc_text}
"""

    if user_prompt:
        prompt += f"\n\n## 用户额外要求\n\n{user_prompt}"

    # 构建消息（支持多模态）
    if _is_multimodal_llm(llm) and images:
        content: List[Dict[str, Any]] = [{"type": "text", "text": prompt}]
        for img_path in images:
            if not os.path.exists(img_path):
                continue
            mime_type, _ = mimetypes.guess_type(img_path)
            if not mime_type:
                mime_type = "image/jpeg"
            base64_str = _encode_image(img_path)
            if not base64_str:
                continue
            data_url = f"data:{mime_type};base64,{base64_str}"
            content.append({"type": "image_url", "image_url": {"url": data_url}})
        messages = [{"role": "user", "content": content}]
    else:
        messages = [{"role": "user", "content": prompt}]

    if DEBUG_LLM:
        print(f"\n{'=' * 60}")
        print("🔍 [DEBUG] LLM 请求 (llm_preprocess_script)")
        print(f"{'=' * 60}")
        print(f"📝 预处理讲稿请求")
        print(f"{'=' * 60}\n")

    response = llm.generate(messages, temperature=0.3)

    if DEBUG_LLM:
        print(f"\n{'=' * 60}")
        print("📥 [DEBUG] LLM 响应 (llm_preprocess_script)")
        print(f"{'=' * 60}")
        print(f"{response[:1000]}...")
        print(f"{'=' * 60}\n")

    # 清理响应：移除可能的 markdown 代码块标记
    result = response.strip()
    if result.startswith("```"):
        # 移除开头的 ```markdown 或 ```
        lines = result.split("\n")
        if lines[0].startswith("```"):
            lines = lines[1:]
        if lines and lines[-1].strip() == "```":
            lines = lines[:-1]
        result = "\n".join(lines)

    return result


def _parse_preprocessed_script(
    preprocessed_text: str,
    image_dir: Path,
) -> List[Dict]:
    """
    解析预处理后的带标记讲稿，返回 blocks 列表。

    Args:
        preprocessed_text: 带【PPT】标记的讲稿文本
        image_dir: 图片目录

    Returns:
        blocks 列表，每个 block 包含 template_hint, text, images
    """
    blocks: List[Dict] = []
    current_block: Optional[Dict] = None

    # 按行解析
    for line in preprocessed_text.split("\n"):
        marker_match = MARKER_RE.match(line.strip())
        if marker_match:
            # 保存之前的 block
            if current_block and current_block.get("text", "").strip():
                blocks.append(current_block)
            # 开始新的 block
            template_num = int(marker_match.group(1))
            current_block = {
                "template_hint": template_num,
                "text": "",
                "images": [],
            }
        elif current_block is not None:
            # 检查是否有图片引用
            img_match = re.search(r"\[图片资源:\s*([^\]]+)\]", line)
            if img_match:
                img_ref = img_match.group(1).strip()
                # 支持完整路径和文件名两种格式
                if os.path.isabs(img_ref) and Path(img_ref).exists():
                    # 完整路径，直接使用
                    current_block["images"].append(img_ref)
                else:
                    # 文件名，拼接 image_dir
                    img_name = Path(img_ref).name  # 取文件名部分
                    img_path = image_dir / img_name
                    if img_path.exists():
                        current_block["images"].append(str(img_path))
                # 从文本中移除图片标记
                line = re.sub(r"\[图片资源:\s*[^\]]+\]", "", line)

            current_block["text"] += line + "\n"

    # 保存最后一个 block
    if current_block and current_block.get("text", "").strip():
        blocks.append(current_block)

    # 清理每个 block 的文本
    for block in blocks:
        block["text"] = block["text"].strip()

    return blocks


def _extract_json_value(text: str, opener: str) -> Any:
    decoder = json.JSONDecoder()
    idx = 0
    while idx < len(text):
        start = text.find(opener, idx)
        if start == -1:
            break
        try:
            value, offset = decoder.raw_decode(text[start:])
            return value
        except json.JSONDecodeError:
            idx = start + 1
    raise ValueError("模型输出中未找到 JSON")


def _ensure_json_object(text: str) -> Dict:
    value = _extract_json_value(text.strip(), "{")
    if not isinstance(value, dict):
        raise ValueError("解析结果不是 JSON 对象")
    return value


def _ensure_json_array(text: str) -> List[Dict]:
    value = _extract_json_value(text.strip(), "[")
    if not isinstance(value, list):
        raise ValueError("解析结果不是 JSON 数组")
    return value


def _coerce_dict(entry):
    if isinstance(entry, dict):
        return entry
    if isinstance(entry, str):
        entry = entry.strip()
        if not entry:
            raise ValueError("模型输出的元素为空字符串，无法解析为 JSON 对象。")
        if entry.startswith("{"):
            try:
                return json.loads(entry)
            except json.JSONDecodeError as exc:
                raise ValueError("模型输出的字符串不是合法 JSON 对象。") from exc
        raise ValueError("字符串元素必须是 JSON 对象字面量。")
    if isinstance(entry, list):
        for candidate in entry:
            try:
                return _coerce_dict(candidate)
            except ValueError:
                continue
        raise ValueError("列表元素中未找到 JSON 对象。")
    raise ValueError("模型输出的元素不是有效的 JSON 对象。")


def choose_llm(
    enable: bool,
    provider: str,
    model: Optional[str],
    base_url: Optional[str] = None,
) -> Optional[BaseLLM]:
    if not enable:
        return None
    provider = (provider or "").lower()

    if provider == "deepseek":
        return DeepSeekLLM(model=model or "deepseek-chat")
    if provider == "local":
        return LocalLLM(model=model)
    if provider == "qwen":
        endpoint = base_url or os.getenv("QWEN_VLLM_BASE_URL")
        if not endpoint:
            raise ValueError(
                "Qwen provider 需要提供 --llm-base-url 或设置 QWEN_VLLM_BASE_URL。"
            )
        return QwenVLLM(base_url=endpoint)
    if provider == "taichu":
        final_model = model or "taichu4_vl_32b"
        return TaichuLLM(model=final_model, base_url=base_url)
    if provider == "glm" or provider == "zhipu":
        final_model = model or "glm-4.5v"
        return GLMLLM(model=final_model, base_url=base_url)
    raise ValueError(f"暂不支持的大模型提供商：{provider}")


def _apply_metadata_overrides(content: Dict, template_info: Dict, metadata: Dict):
    if not metadata:
        return
    for field in template_info["text_fields"]:
        path = list(field["path"])
        key = "/".join(path)
        value = None
        if metadata.get("lecturer") and "主讲" in key:
            value = metadata["lecturer"]
        elif metadata.get("college") and "学院" in key:
            value = metadata["college"]
        elif metadata.get("course") and any(
            token in key for token in ("课程", "课程名称", "项目")
        ):
            value = metadata["course"]
        if value is not None:
            _assign_in_schema(content, path, value)


def _apply_section_context(
    content: Dict,
    template_info: Dict,
    section_context: Dict[str, Optional[str]],
    title_hint: Optional[str] = None,
):
    """
    根据章节上下文填充章节/知识点相关字段。

    如果 LLM 没有正确填充这些字段，通过关键词匹配进行备用填充。
    """
    if not section_context:
        return

    # 字段名关键词 → 章节上下文 key 的映射
    field_mappings = [
        # 一级标题（章节）
        (["章节", "一级标题", "章节名", "大章节"], "chapter"),
        # 二级标题（知识点）
        (["知识点", "二级标题", "小节名", "节名"], "section"),
        # 三级标题
        (["三级标题", "小节", "子节"], "subsection"),
    ]

    def _trim_to_max(value: str, max_chars: Optional[int]) -> str:
        if max_chars and len(value) > max_chars:
            return value[:max_chars]
        return value
    def _norm(text: Optional[str]) -> str:
        return re.sub(r"[（()）\s]", "", text or "")

    for field in template_info["text_fields"]:
        path = list(field["path"])
        key = "/".join(path)
        max_chars = field.get("max_chars")

        current_value = _get_in_schema(content, path) or ""
        required = bool(field.get("required"))

        # 根据字段名匹配章节上下文
        for keywords, ctx_key in field_mappings:
            if any(kw in key for kw in keywords):
                target = section_context.get(ctx_key)

                # 为三级标题增加回退：仅在字段必填时才用当前页首句摘要
                if ctx_key == "subsection" and not target and required and title_hint:
                    target = title_hint

                if not target:
                    break

                target = _trim_to_max(str(target), max_chars)

                # 必填字段：如果与上下文不同则覆盖；可选字段：仅在为空时填充
                if required:
                    if _norm(current_value) != _norm(target):
                        _assign_in_schema(content, path, target)
                else:
                    if not current_value:
                        _assign_in_schema(content, path, target)
                break


def _get_in_schema(content: Dict, path: List[str]) -> Optional[str]:
    """从嵌套字典中获取值。"""
    node = content
    for key in path:
        if isinstance(node, dict) and key in node:
            node = node[key]
        else:
            return None
    # 处理 {"type": "text", "value": "xxx"} 格式
    if isinstance(node, dict) and "value" in node:
        return node.get("value")
    return node if isinstance(node, str) else None


def _fill_with_template(
    template_num: int,
    template_info: Dict,
    block: Dict,
    llm: Optional[BaseLLM],
    metadata: Dict,
    user_prompt: Optional[str] = None,
    use_multimodal: bool = True,
    section_context: Optional[Dict[str, Optional[str]]] = None,
) -> Dict:
    """
    使用模板填充单个 block 的内容。

    Args:
        use_multimodal: 是否使用多模态消息（默认 True）
                       当讲稿有 PPT 标记时，图片位置已确定，建议设为 False
        section_context: 章节上下文（chapter/section/subsection）
    """
    content = llm_fill_slide(
        llm,
        template_info,
        block.get("text", ""),
        block.get("images", []),
        user_prompt,
        use_multimodal,
        metadata,  # 传递 metadata 给 LLM
        section_context,  # 传递章节上下文给 LLM
    )
    # 备用：如果 LLM 没有正确填充 metadata，尝试通过关键词匹配填充
    _apply_metadata_overrides(content, template_info, metadata)
    return {
        "page_type": template_info["page_type"],
        "template_page_num": template_num,
        "content": content,
    }


def _strip_values(node):
    if isinstance(node, dict):
        if "type" in node and "value" in node:
            return node.get("value", "")
        return {k: _strip_values(v) for k, v in node.items()}
    if isinstance(node, list):
        return [_strip_values(item) for item in node]
    return node


def _empty_content(template_info: Dict) -> Dict:
    result = _clone_schema(template_info["schema"])
    for field in template_info["text_fields"]:
        _assign_in_schema(result, list(field["path"]), "")
    for field in template_info["image_fields"]:
        _assign_in_schema(result, list(field["path"]), "")
    return result


def _prepend_cover_page(
    pages: List[Dict],
    templates: Dict[int, Dict],
    metadata: Dict = None,
    llm: Optional[BaseLLM] = None,
):
    """在页面列表开头插入封面页，并应用元数据（课程名称、学院名称、讲师名称）。"""
    cover_template = templates.get(1)
    if not cover_template:
        return
    if pages and pages[0].get("template_page_num") == 1:
        # 封面页已存在，确保应用 metadata
        if metadata:
            _apply_metadata_overrides(
                pages[0].get("content", {}), cover_template, metadata
            )
        return

    # 创建封面页内容
    if llm and metadata:
        # 使用 LLM 智能填充封面页（让 LLM 根据字段名称语义匹配 metadata）
        cover_text = "这是封面页，请根据预设元信息填充相应字段。"
        content = llm_fill_slide(
            llm,
            cover_template,
            cover_text,
            [],  # 封面页通常没有图片
            None,  # user_prompt
            False,  # use_multimodal
            metadata,
        )
    else:
        # 没有 LLM，使用关键词匹配
        content = _empty_content(cover_template)
        if metadata:
            _apply_metadata_overrides(content, cover_template, metadata)

    pages.insert(
        0,
        {
            "page_type": cover_template["page_type"],
            "template_page_num": 1,
            "content": content,
        },
    )


def _fill_by_markers(
    blocks: List[Dict],
    templates: Dict[int, Dict],
    llm: Optional[BaseLLM],
    metadata: Dict,
    user_prompt: Optional[str] = None,
) -> List[Dict]:
    """
    按照讲稿中的 PPT 标记填充内容。

    由于讲稿已有明确的标记，图片位置已经确定（每个 block 的 images 字段），
    因此不需要使用多模态模型来界定图片位置，设置 use_multimodal=False。

    新增：跟踪章节上下文，让 LLM 知道当前内容属于哪个章节/知识点。
    """
    pages: List[Dict] = []
    section_ctx = SectionContext()  # 章节上下文跟踪器

    for block in blocks:
        template_num = block.get("template_hint")
        if template_num is None:
            continue
        if template_num not in templates:
            raise ValueError(
                f"模板 {template_num} 未在 template.json 中定义或不在 template.txt 中允许。"
            )

        template_info = templates[template_num]
        page_type = template_info.get("page_type", "")
        block_text = block.get("text", "")
        lines = [l.strip() for l in block_text.split("\n") if l.strip()]

        # 对于章节页，特殊处理：第一行是章节名，第二行是知识点名
        if "章节" in page_type and lines:
            # 第一行作为章节名称（一级标题）
            section_ctx.level1 = lines[0]
            section_ctx.level2 = None
            section_ctx.level3 = None
            # 第二行如果存在，作为知识点名称（二级标题）
            if len(lines) > 1:
                section_ctx.level2 = lines[1]
        else:
            # 非章节页：检查文本的前几行是否包含章节标题格式
            for line in lines[:5]:
                section_ctx.update(line)

        pages.append(
            _fill_with_template(
                template_num,
                template_info,
                block,
                llm,
                metadata,
                user_prompt,
                use_multimodal=False,  # 有标记时图片位置已确定，不需要多模态
                section_context=section_ctx.to_dict(),  # 传递章节上下文
            )
        )
    return pages


def _preprocess_and_fill(
    blocks: List[Dict],
    templates: Dict[int, Dict],
    llm: Optional[BaseLLM],
    metadata: Dict,
    user_prompt: Optional[str] = None,
    run_dir: Optional[Path] = None,
    has_marker: bool = False,
) -> List[Dict]:
    """
    统一的讲稿处理流程（两步处理）。

    无论讲稿是否已有【PPT】标记，都统一走此流程：
    1. 预处理：让 LLM 优化语法；如果已有标记则保持分页，否则自动分页
    2. 填充：根据模板字段精简内容，生成最终 JSON

    Args:
        blocks: 原始讲稿的 block 列表
        templates: 模板定义字典
        llm: LLM 实例
        metadata: 元数据
        user_prompt: 用户自定义提示
        run_dir: 运行目录，用于保存中间讲稿
        has_marker: 原始讲稿是否已有【PPT】标记

    Returns:
        填充后的页面列表
    """
    if not llm:
        raise ValueError("未启用 LLM，无法处理讲稿。")

    # 合并所有 block 的文本和图片，使用 block["images"] 中的完整路径
    if has_marker:
        # 已有标记：保留【PPT】标记结构
        doc_parts = []
        for block in blocks:
            text = block.get("text", "")
            template_hint = block.get("template_hint")
            # 构建页面内容（文本 + 图片引用）
            page_content = text
            for img_path in block.get("images", []):
                page_content += f"\n[图片资源: {img_path}]"
            if template_hint is not None:
                doc_parts.append(f"【PPT{template_hint}】\n{page_content}")
            elif page_content.strip():
                doc_parts.append(page_content)
        doc_text = "\n\n".join(doc_parts)
    else:
        # 无标记：合并所有文本和图片
        doc_parts = []
        for block in blocks:
            text = block.get("text", "")
            for img_path in block.get("images", []):
                text += f"\n[图片资源: {img_path}]"
            if text.strip():
                doc_parts.append(text)
        doc_text = "\n\n".join(doc_parts)
    all_images = [path for block in blocks for path in block.get("images", [])]

    # Step 1: 预处理讲稿
    if has_marker:
        print("📝 Step 1: 优化讲稿文本（保持原有分页）...")
    else:
        print("📝 Step 1: 预处理讲稿（自动分页）...")
    preprocessed_script = llm_preprocess_script(
        llm, doc_text, templates, all_images, user_prompt, has_marker
    )

    # 保存中间讲稿到文件（供管理员/开发者下载）
    if run_dir:
        script_path = run_dir / "preprocessed_script.md"
        script_path.write_text(preprocessed_script, encoding="utf-8")
        print(f"💾 中间讲稿已保存: {script_path}")

    # Step 2: 解析中间讲稿为 blocks
    image_dir = run_dir / "images" if run_dir else Path(".")
    preprocessed_blocks = _parse_preprocessed_script(preprocessed_script, image_dir)

    if not preprocessed_blocks:
        raise ValueError("预处理后的讲稿没有有效的【PPT】标记，请检查 LLM 输出。")

    print(f"✅ 预处理完成，共 {len(preprocessed_blocks)} 个页面")

    # Step 3: 复用 _fill_by_markers 处理
    print("📝 Step 2: 填充页面内容...")
    pages = _fill_by_markers(preprocessed_blocks, templates, llm, metadata, user_prompt)

    return pages


def generate_config_data(
    docx_path: str,
    template_json: str,
    template_list: str,
    use_llm: bool,
    llm_provider: str,
    llm_model: Optional[str],
    llm_base_url: Optional[str],
    metadata_overrides: Optional[Dict[str, str]],
    run_dir: Path,
    user_prompt: Optional[str] = None,
) -> Dict:
    """核心逻辑：生成 JSON 内容，供 GUI/CLI 复用。"""
    metadata_overrides = metadata_overrides or {}
    image_dir = run_dir / "images"
    blocks, has_marker, metadata = parse_docx_blocks(docx_path, image_dir)
    for key in ("course", "college", "lecturer"):
        if metadata_overrides.get(key):
            metadata[key] = metadata_overrides[key]

    templates = load_template_defs(template_json, template_list)
    llm = choose_llm(use_llm, llm_provider, llm_model, llm_base_url)

    # 统一走预处理流程：
    # - 如果已有标记：保持分页，只优化文本
    # - 如果无标记：自动分页 + 优化文本
    pages = _preprocess_and_fill(
        blocks, templates, llm, metadata, user_prompt, run_dir, has_marker
    )

    if not pages:
        raise ValueError("未生成任何幻灯片内容，请检查讲稿或模板。")

    _prepend_cover_page(pages, templates, metadata, llm)

    stripped_pages = []
    for page in pages:
        stripped_pages.append(
            {
                "page_type": page.get("page_type"),
                "template_page_num": page.get("template_page_num"),
                "content": _strip_values(page.get("content", {})),
            }
        )

    return {"ppt_pages": stripped_pages}


def process_docx(
    docx_path: str,
    template_json: str,
    template_list: str,
    output_path: Optional[str],
    use_llm: bool,
    llm_provider: str,
    llm_model: Optional[str],
    llm_base_url: Optional[str],
    override_course: Optional[str],
    override_college: Optional[str],
    override_lecturer: Optional[str],
    run_dir: Optional[str],
    config_name: str,
):
    """CLI 包装：处理参数、保证 run 目录存在，并额外复制文件到 output。"""
    metadata_overrides = {
        "course": override_course,
        "college": override_college,
        "lecturer": override_lecturer,
    }

    base_dir = Path(run_dir) if run_dir else _create_run_dir()
    base_dir.mkdir(parents=True, exist_ok=True)
    config_path = base_dir / config_name

    config = generate_config_data(
        docx_path,
        template_json,
        template_list,
        use_llm,
        llm_provider,
        llm_model,
        llm_base_url,
        metadata_overrides,
        base_dir,
    )

    config_path.write_text(
        json.dumps(config, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    if output_path:
        explicit = Path(output_path)
        explicit.parent.mkdir(parents=True, exist_ok=True)
        shutil.copyfile(config_path, explicit)
        print(f"📄 另存为：{explicit}")

    print(f"✅ 已生成 JSON：{config_path}")
    print(f"📁 资源输出目录：{base_dir}")


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="根据 DOCX 讲稿生成 PPT 配置 JSON。")
    parser.add_argument("--docx", default=None, help="讲稿 DOCX 路径")
    parser.add_argument(
        "--from-preprocessed",
        default=None,
        help="从预处理后的讲稿(.md)继续，跳过预处理步骤",
    )
    parser.add_argument(
        "--template-json", default="template/template.json", help="模板定义 JSON 文件"
    )
    parser.add_argument(
        "--template-list", default="template/template.txt", help="可用模板编号列表 txt"
    )
    parser.add_argument(
        "--output",
        default=None,
        help="如需额外复制一份 JSON，请提供完整路径；若省略则仅在 temp/run-*/ 中生成",
    )
    parser.add_argument("--use-llm", action="store_true", help="启用大模型填充/排版")
    parser.add_argument("--llm-provider", default="deepseek", help="大模型提供商")
    parser.add_argument("--llm-model", default="deepseek-chat", help="大模型名称")
    parser.add_argument(
        "--llm-base-url",
        default="http://172.18.75.58:9000",
        help="自定义大模型接口地址",
    )
    parser.add_argument("--course-name", default=None, help="手动指定课程/项目名称")
    parser.add_argument("--college-name", default=None, help="手动指定学院/单位")
    parser.add_argument("--lecturer-name", default=None, help="手动指定主讲教师姓名")
    parser.add_argument(
        "--run-dir",
        default=None,
        help="指定输出目录（默认在 temp 下自动创建 run-时间戳-随机值 文件夹）",
    )
    parser.add_argument(
        "--config-name",
        default="config.json",
        help="输出目录中生成的配置文件名称",
    )
    parser.add_argument(
        "--template-pptx",
        default=None,
        help="模板 PPTX 路径（用于 --from-preprocessed 模式）",
    )
    return parser


def continue_from_preprocessed(
    preprocessed_path: str,
    template_json: str,
    template_pptx: str,
    use_llm: bool = True,
    llm_provider: str = "deepseek",
    llm_model: Optional[str] = None,
    llm_base_url: Optional[str] = None,
    override_course: Optional[str] = None,
    override_college: Optional[str] = None,
    override_lecturer: Optional[str] = None,
    run_dir: Optional[str] = None,
    config_name: str = "config.json",
) -> Dict:
    """
    从预处理后的讲稿继续生成 config.json。

    Args:
        preprocessed_path: 预处理后的讲稿(.md)路径
        template_json: 模板定义 JSON 文件路径
        template_pptx: 模板 PPTX 文件路径
        其他参数与 process_docx 相同

    Returns:
        生成的 config 数据
    """
    preprocessed_file = Path(preprocessed_path)
    if not preprocessed_file.exists():
        raise FileNotFoundError(f"预处理讲稿不存在: {preprocessed_path}")

    # 确定 run_dir（默认使用预处理文件所在目录）
    if run_dir:
        run_dir_path = Path(run_dir)
    else:
        run_dir_path = preprocessed_file.parent

    image_dir = run_dir_path / "images"

    # 读取预处理后的讲稿
    preprocessed_text = preprocessed_file.read_text(encoding="utf-8")

    # 解析成 blocks
    blocks = _parse_preprocessed_script(preprocessed_text, image_dir)
    print(f"✅ 解析到 {len(blocks)} 个 blocks")

    # 加载模板
    templates = load_template_defs(template_json, None)
    print(f"✅ 加载了 {len(templates)} 个模板页")

    # 初始化 LLM
    llm = choose_llm(use_llm, llm_provider, llm_model, llm_base_url)

    # 准备 metadata
    metadata = {
        "course": override_course or "",
        "college": override_college or "",
        "lecturer": override_lecturer or "",
    }

    # 填充内容
    print("🔄 开始填充内容...")
    pages = _fill_by_markers(blocks, templates, llm, metadata, user_prompt=None)
    print(f"✅ 生成了 {len(pages)} 页")

    # 构建 config
    config = {
        "template": template_pptx,
        "pages": pages,
    }

    # 保存
    config_path = run_dir_path / config_name
    config_path.write_text(
        json.dumps(config, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(f"✅ 已保存到 {config_path}")

    return config


def main():
    args = build_arg_parser().parse_args()

    # 从预处理讲稿继续
    if args.from_preprocessed:
        if not args.template_pptx:
            print("❌ 使用 --from-preprocessed 时必须指定 --template-pptx")
            return
        continue_from_preprocessed(
            preprocessed_path=args.from_preprocessed,
            template_json=args.template_json,
            template_pptx=args.template_pptx,
            use_llm=args.use_llm,
            llm_provider=args.llm_provider,
            llm_model=args.llm_model,
            llm_base_url=args.llm_base_url,
            override_course=args.course_name,
            override_college=args.college_name,
            override_lecturer=args.lecturer_name,
            run_dir=args.run_dir,
            config_name=args.config_name,
        )
        return

    # 正常流程：从 DOCX 开始
    if not args.docx:
        print("❌ 必须指定 --docx 或 --from-preprocessed")
        return

    process_docx(
        docx_path=args.docx,
        template_json=args.template_json,
        template_list=args.template_list,
        output_path=args.output,
        use_llm=args.use_llm,
        llm_provider=args.llm_provider,
        llm_model=args.llm_model,
        llm_base_url=args.llm_base_url,
        override_course=args.course_name,
        override_college=args.college_name,
        override_lecturer=args.lecturer_name,
        run_dir=args.run_dir,
        config_name=args.config_name,
    )


def _strip_values(node):
    if isinstance(node, dict):
        if "type" in node and "value" in node:
            return node.get("value", "")
        return {k: _strip_values(v) for k, v in node.items()}
    if isinstance(node, list):
        return [_strip_values(item) for item in node]
    return node


if __name__ == "__main__":
    main()
