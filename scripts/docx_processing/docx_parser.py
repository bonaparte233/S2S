"""DOCX 文件解析模块。"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.oxml.ns import qn

from scripts.docx_processing.constants import (
    MARKER_RE,
    IMAGE_NAME_TEMPLATE,
    HEADING_L1_RE,
    HEADING_L2_RE,
)


def _extract_heading_level(text: str) -> Tuple[Optional[int], Optional[str]]:
    """从文本中提取章节标题级别和标题内容。

    Returns:
        (level, title): level 为 1/2 表示一/二级标题，None 表示不是标题
    """
    text = text.strip()
    if not text:
        return None, None

    if HEADING_L1_RE.match(text):
        return 1, text
    if HEADING_L2_RE.match(text):
        return 2, text

    return None, None


class SectionContext:
    """章节上下文，跟踪当前所处的章节层级（只追踪两级）。"""

    def __init__(self):
        self.level1: Optional[str] = None  # 章节
        self.level2: Optional[str] = None  # 知识点

    def update(self, text: str) -> None:
        """根据文本内容更新章节上下文。"""
        level, title = _extract_heading_level(text)
        if level == 1:
            self.level1 = title
            self.level2 = None
        elif level == 2:
            self.level2 = title

    def to_dict(self) -> Dict[str, Optional[str]]:
        """返回当前章节上下文的字典表示。"""
        return {
            "chapter": self.level1,
            "section": self.level2,
        }

    def __str__(self) -> str:
        parts = []
        if self.level1:
            parts.append(f"章节：{self.level1}")
        if self.level2:
            parts.append(f"知识点：{self.level2}")
        return " > ".join(parts) if parts else "（无章节信息）"


def _extract_images_from_paragraph(
    paragraph, image_dir: Path, image_counter: int
) -> Tuple[List[str], int]:
    """从段落中提取图片。

    Returns:
        (images, new_counter): 图片路径列表和更新后的计数器
    """
    images = []
    for element in paragraph._p.iter():
        if element.tag not in {qn("a:blip"), qn("pic:blip")}:
            continue
        r_id = element.get(qn("r:embed"))
        if not r_id:
            continue
        part = paragraph.part.related_parts.get(r_id)
        if not part:
            continue
        ext = part.filename.split(".")[-1].lower() or "png"
        filename = image_dir / IMAGE_NAME_TEMPLATE.format(idx=image_counter, ext=ext)
        with open(filename, "wb") as f:
            f.write(part.blob)
        images.append(str(filename))
        image_counter += 1
    return images, image_counter


def _extract_metadata(text: str) -> Optional[Tuple[str, str]]:
    """从文本中提取元数据（课程名称、学院名称、主讲教师）。

    Returns:
        (key, value) 或 None
    """
    match = re.match(r"^(课程名称|学院名称|主讲教师)[：:]\s*(.+)$", text.strip())
    if match:
        return match.group(1), match.group(2).strip()
    return None


def parse_docx_blocks(
    doc_path: str, image_dir: Path
) -> Tuple[List[Dict], bool, Dict]:
    """读取 DOCX 并按照 PPT 标记拆分内容。

    Args:
        doc_path: DOCX 文件路径
        image_dir: 图片保存目录

    Returns:
        (slides, has_marker, metadata):
        - slides: 内容块列表
        - has_marker: 是否包含 PPT 标记
        - metadata: 元数据字典
    """
    doc = Document(doc_path)
    image_dir.mkdir(parents=True, exist_ok=True)

    slides: List[Dict] = []
    current: Optional[Dict] = None
    buffer: List[str] = []
    has_marker = False
    image_counter = 1
    metadata: Dict[str, str] = {}

    def flush():
        nonlocal current, buffer
        if current is None and buffer:
            slides.append({
                "template_hint": None,
                "text": "\n".join(buffer).strip(),
                "images": [],
            })
        elif current:
            current["text"] = "\n".join(buffer).strip()
            slides.append(current)
        buffer = []
        current = None

    def ensure_block():
        nonlocal current
        if current is None:
            current = {"template_hint": None, "text": "", "images": []}

    def attach_images(paths: List[str]):
        if not paths:
            return
        ensure_block()
        current.setdefault("images", []).extend(paths)

    # 处理段落
    for para in doc.paragraphs:
        text = para.text
        stripped = text.strip()

        # 提取元数据
        meta_result = _extract_metadata(stripped)
        if meta_result:
            key, value = meta_result
            if key == "课程名称":
                metadata["course"] = value
            elif key == "学院名称":
                metadata["college"] = value
            elif key == "主讲教师":
                metadata["lecturer"] = value
            continue

        # 提取图片
        images, image_counter = _extract_images_from_paragraph(
            para, image_dir, image_counter
        )
        if images:
            attach_images(images)

        # 处理 PPT 标记
        matches = list(MARKER_RE.finditer(text))
        if matches:
            idx = 0
            for match in matches:
                prefix = text[idx:match.start()].strip()
                if prefix:
                    buffer.append(prefix)
                flush()
                has_marker = True
                current = {
                    "template_hint": int(match.group(1)),
                    "text": "",
                    "images": [],
                }
                idx = match.end()
            remainder = text[idx:].strip()
            if remainder:
                buffer.append(remainder)
            continue

        if stripped:
            buffer.append(stripped)

    # 提取表格内容
    if doc.tables:
        from scripts.docx_processing.docx_table_parser import extract_tables_from_docx
        table_results = extract_tables_from_docx(doc)
        for table_result in table_results:
            if table_result.raw_text:
                buffer.append("\n--- 表格内容 ---\n")
                buffer.append(table_result.raw_text)

    flush()

    if not slides:
        full_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        slides.append({"template_hint": None, "text": full_text, "images": []})

    return slides, has_marker, metadata
