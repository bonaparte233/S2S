"""æ ¹æ® DOCX è®²ç¨¿å’Œæ¨¡æ¿å®šä¹‰ç”Ÿæˆ JSONï¼Œå¯é€‰è°ƒç”¨ DeepSeek LLMã€‚"""

from __future__ import annotations

import argparse
import json
import re
import secrets
import shutil
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.oxml.ns import qn

from llm_client import BaseLLM, DeepSeekLLM, LocalLLM

MARKER_RE = re.compile(r"ã€PPT(\d+)ã€‘")
IMAGE_NAME_TEMPLATE = "doc_image_{idx}.{ext}"


def _create_run_dir(base_dir: Path = Path("temp")) -> Path:
    """åˆ›å»ºå¸¦æ—¶é—´æˆ³å‰ç¼€çš„è¿è¡Œç›®å½•ï¼Œæ–¹ä¾¿å‰ç«¯ä¸€æ¬¡å¤„ç†å¯¹åº”åˆ°å•ç‹¬ç›®å½•ã€‚"""
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    suffix = secrets.token_hex(2)
    run_dir = base_dir / f"script-{timestamp}-{suffix}"
    run_dir.mkdir(parents=True, exist_ok=True)
    return run_dir


def parse_docx_blocks(doc_path: str, image_dir: Path) -> Tuple[List[Dict], bool, Dict]:
    """è¯»å– DOCX å¹¶æŒ‰ç…§ PPT æ ‡è®°æ‹†åˆ†å†…å®¹ï¼ŒåŒæ—¶ä¿å­˜æå–çš„å›¾ç‰‡ä¸è¯¾ç¨‹å…ƒä¿¡æ¯ã€‚"""
    doc = Document(doc_path)
    image_dir.mkdir(parents=True, exist_ok=True)
    slides: List[Dict] = []
    current: Optional[Dict] = None
    buffer: List[str] = []
    has_marker = False
    image_counter = 1
    metadata: Dict[str, str] = {}

    def flush():
        nonlocal current, buffer
        if current is None and buffer:
            slides.append(
                {"template_hint": None, "text": "\n".join(buffer).strip(), "images": []}
            )
        elif current:
            current["text"] = "\n".join(buffer).strip()
            slides.append(current)
        buffer = []
        current = None

    def ensure_block():
        nonlocal current
        if current is None:
            current = {"template_hint": None, "text": "", "images": []}

    def attach_images(paths: List[str]):
        if not paths:
            return
        ensure_block()
        current.setdefault("images", []).extend(paths)

    def extract_images(paragraph) -> List[str]:
        nonlocal image_counter
        images = []
        for element in paragraph._p.iter():
            if element.tag not in {qn("a:blip"), qn("pic:blip")}:
                continue
            r_id = element.get(qn("r:embed"))
            if not r_id:
                continue
            part = paragraph.part.related_parts.get(r_id)
            if not part:
                continue
            ext = part.filename.split(".")[-1].lower() or "png"
            filename = image_dir / IMAGE_NAME_TEMPLATE.format(
                idx=image_counter, ext=ext
            )
            with open(filename, "wb") as f:
                f.write(part.blob)
            images.append(str(filename))
            image_counter += 1
        return images

    for para in doc.paragraphs:
        text = para.text
        stripped = text.strip()
        meta_match = re.match(r"^(è¯¾ç¨‹åç§°|å­¦é™¢åç§°|ä¸»è®²æ•™å¸ˆ)[ï¼š:]\s*(.+)$", stripped)
        if meta_match:
            key = meta_match.group(1)
            value = meta_match.group(2).strip()
            if key == "è¯¾ç¨‹åç§°":
                metadata["course"] = value
            elif key == "å­¦é™¢åç§°":
                metadata["college"] = value
            elif key == "ä¸»è®²æ•™å¸ˆ":
                metadata["lecturer"] = value
            continue
        images = extract_images(para)
        if images:
            attach_images(images)

        matches = list(MARKER_RE.finditer(text))
        if matches:
            idx = 0
            for match in matches:
                prefix = text[idx : match.start()].strip()
                if prefix:
                    buffer.append(prefix)
                flush()
                has_marker = True
                current = {
                    "template_hint": int(match.group(1)),
                    "text": "",
                    "images": [],
                }
                idx = match.end()
            remainder = text[idx:].strip()
            if remainder:
                buffer.append(remainder)
            continue

        if stripped:
            buffer.append(stripped)

    flush()
    if not slides:
        full_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        slides.append({"template_hint": None, "text": full_text, "images": []})
    return slides, has_marker, metadata


def load_template_defs(template_json: str, template_list: str) -> Dict[int, Dict]:
    data = json.loads(Path(template_json).read_text(encoding="utf-8"))
    allowed = None
    if template_list and Path(template_list).exists():
        allowed = {
            int(item.strip())
            for item in Path(template_list)
            .read_text(encoding="utf-8")
            .replace(",", " ")
            .split()
            if item.strip().isdigit()
        }
    templates = {}
    manifest = {
        item["template_page_num"]: item
        for item in data.get("manifest", [])
        if "template_page_num" in item
    }
    for page in data.get("ppt_pages", []):
        num = page.get("template_page_num")
        if not isinstance(num, int):
            continue
        if allowed and num not in allowed:
            continue
        schema = page.get("content", {})
        fields = _collect_fields(schema)
        templates[num] = {
            "page_type": page.get("page_type", f"æ¨¡æ¿ç¬¬{num}é¡µ"),
            "schema": schema,
            "meta": page.get("meta", {}) or manifest.get(num, {}),
            "text_fields": [field for field in fields if not field["is_image"]],
            "image_fields": [field for field in fields if field["is_image"]],
        }
    if not templates:
        raise ValueError("æ¨¡æ¿åˆ—è¡¨ä¸ºç©ºï¼Œæ— æ³•åŒ¹é…ã€‚")
    return templates


def _collect_fields(schema, prefix=None):
    results = []
    prefix = prefix or ()
    if isinstance(schema, dict):
        if "type" in schema and "value" in schema:
            field_type = schema.get("type") or "text"
            results.append(
                {
                    "path": prefix,
                    "is_image": field_type.lower() == "image",
                    "hint": schema.get("hint") or "",
                    "max_chars": schema.get("max_chars"),
                    "required": bool(schema.get("required", False)),
                }
            )
        else:
            for key, value in schema.items():
                results.extend(_collect_fields(value, prefix + (key,)))
    else:
        # å…¼å®¹æ—§æ ¼å¼ï¼šå­—ç¬¦ä¸²æˆ–å…¶ä»–ç±»å‹è§†ä¸ºæ–‡æœ¬å¶å­
        results.append(
            {
                "path": prefix,
                "is_image": any(
                    "å›¾ç‰‡" in seg or "image" in seg.lower() for seg in prefix
                ),
                "hint": "",
                "max_chars": None,
                "required": False,
            }
        )
    return results


def _clone_schema(schema: Dict) -> Dict:
    return json.loads(json.dumps(schema, ensure_ascii=False))


def _assign_in_schema(schema: Dict, path: List[str], value: str):
    node = schema
    for key in path[:-1]:
        node = node.setdefault(key, {})
    leaf = node.get(path[-1])
    if isinstance(leaf, dict) and "type" in leaf:
        leaf["value"] = value
    else:
        node[path[-1]] = value


def _simple_fill(template_info: Dict, raw_text: str, images: List[str]) -> Dict:
    result = _clone_schema(template_info["schema"])
    text_fields = template_info["text_fields"]
    image_fields = template_info["image_fields"]
    if text_fields:
        _assign_in_schema(result, list(text_fields[0]["path"]), raw_text)
        for field in text_fields[1:]:
            _assign_in_schema(result, list(field["path"]), "")
    for idx, field in enumerate(image_fields):
        value = images[idx] if idx < len(images) else ""
        _assign_in_schema(result, list(field["path"]), value)
    return result


def _build_prompt(template_info: Dict, raw_text: str, images: List[str]) -> str:
    def describe_fields(fields):
        if not fields:
            return "æ— "
        lines = []
        for idx, field in enumerate(fields, 1):
            name = "/".join(field["path"])
            hint = field.get("hint") or "å¡«å†™å†…å®¹"
            extra = []
            if field.get("max_chars"):
                extra.append(f"â‰¤{field['max_chars']}å­—")
            if field.get("required"):
                extra.append("å¿…å¡«")
            extra_note = f"ï¼ˆ{'ï¼Œ'.join(extra)}ï¼‰" if extra else ""
            lines.append(f"{idx}. {name}ï¼š{hint}{extra_note}")
        return "\n".join(lines)

    text_desc = describe_fields(template_info["text_fields"])
    image_desc = describe_fields(template_info["image_fields"])
    image_section = "æ— " if not images else "\n".join(images)
    meta = template_info.get("meta") or {}
    scene = "ã€".join(meta.get("scene", [])) or "é€šç”¨"
    layout = meta.get("layout", template_info["page_type"])
    style = meta.get("style", "")
    note = meta.get("notes", "")
    prompt = f"""
è¯·é˜…è¯»ä»¥ä¸‹è®²ç¨¿å¹¶ç”Ÿæˆä¸€ä¸ª JSONï¼Œå¯¹æ¨¡æ¿ã€Š{template_info["page_type"]}ã€‹çš„æ–‡æœ¬/å›¾ç‰‡å­—æ®µè¿›è¡Œå¡«å……ã€‚
æ¨¡æ¿å¸ƒå±€ï¼š{layout}ï¼›ä½¿ç”¨åœºæ™¯ï¼š{scene}ï¼›é£æ ¼æç¤ºï¼š{style}
æ³¨æ„äº‹é¡¹ï¼š{note}
åŠ¡å¿…è®°ä½è®²ç¨¿ä¸­æåˆ°çš„ä¸»è®²äººå§“åã€è¯¾ç¨‹/è®²åº§/é¡¹ç›®åç§°æˆ–å…¶ä»–å…³é”®ä¸“æœ‰åè¯ï¼Œå¹¶åœ¨åç»­æ‰€æœ‰éœ€è¦è¿™äº›ä¿¡æ¯çš„å­—æ®µä¿æŒå®Œå…¨ä¸€è‡´ã€ä¸è¦æ”¹å†™ã€‚æ‰€æœ‰æ ‡è®°ä¸ºâ€œrequiredâ€çš„å­—æ®µå¿…é¡»å¡«å†™ï¼Œä¸”æ–‡æœ¬é•¿åº¦ä¸å¾—è¶…è¿‡å¯¹åº”çš„ max_chars é™åˆ¶ã€‚
è¯¥æ¨¡æ¿åŒ…å«å¦‚ä¸‹æ–‡æœ¬å­—æ®µï¼ˆæŒ‰ç…§é¡ºåºå¯¹åº”ï¼‰ï¼š
{text_desc}

å›¾ç‰‡å­—æ®µï¼ˆè‹¥æ— å¯ç•™ç©ºï¼‰ï¼š
{image_desc}

å¯ç”¨å›¾ç‰‡è·¯å¾„ï¼š
{image_section}

è¾“å‡ºæ ¼å¼ç¤ºä¾‹ï¼š
{{
  "texts": ["æ–‡æœ¬1", "æ–‡æœ¬2", "..."],
  "images": ["å›¾ç‰‡è·¯å¾„1", "å›¾ç‰‡è·¯å¾„2", "..."]
}}

è¯·ä¸¥æ ¼ä¿æŒæ•°ç»„é•¿åº¦ä¸å­—æ®µæ•°é‡ä¸€è‡´ï¼Œtexts[1] å¿…é¡»å¯¹åº”ä¸Šè¿°åˆ—è¡¨ä¸­çš„ç¬¬ 1 ä¸ªæ–‡æœ¬å­—æ®µï¼Œä¾æ­¤ç±»æ¨ã€‚
è®²ç¨¿å†…å®¹ï¼š
{raw_text}
"""
    return prompt


def _lookup_field_value(field, payload, fallback_list, idx):
    path = list(field["path"])
    key = "/".join(path)

    def normalize(text):
        return "".join(ch for ch in text.lower() if ch not in {" ", "_"})

    if isinstance(payload, dict):
        norm_targets = {normalize(key), normalize(path[-1])}
        for candidate_key, candidate_val in payload.items():
            norm = normalize(candidate_key)
            if norm in norm_targets or any(
                target and target in norm for target in norm_targets
            ):
                return candidate_val
        values = list(payload.values())
        if idx < len(values):
            return values[idx]
        return values[-1] if values else ""

    if isinstance(fallback_list, list):
        return fallback_list[idx] if idx < len(fallback_list) else ""
    return ""


def llm_fill_slide(
    llm: BaseLLM, template_info: Dict, raw_text: str, images: List[str]
) -> Dict:
    if not llm:
        return _simple_fill(template_info, raw_text, images)

    prompt = _build_prompt(template_info, raw_text, images)
    response = llm.generate([{"role": "user", "content": prompt}], temperature=0.2)
    try:
        data = _ensure_json_object(response)
        texts = data.get("texts", [])
        imgs = data.get("images", [])
    except Exception:
        texts, imgs = [raw_text], images

    result = _clone_schema(template_info["schema"])
    text_fields = template_info["text_fields"]
    image_fields = template_info["image_fields"]

    for idx, field in enumerate(text_fields):
        value = _lookup_field_value(
            field, texts, texts if isinstance(texts, list) else None, idx
        )
        _assign_in_schema(result, list(field["path"]), value)

    for idx, field in enumerate(image_fields):
        value = _lookup_field_value(
            field, imgs, imgs if isinstance(imgs, list) else None, idx
        )
        if not value and isinstance(images, list) and idx < len(images):
            value = images[idx]
        _assign_in_schema(result, list(field["path"]), value)

    return result


def llm_plan_slides(
    llm: BaseLLM, doc_text: str, templates: Dict[int, Dict], images: List[str]
) -> List[Dict]:
    template_desc = "\n".join(
        f"- æ¨¡æ¿ {info['page_type']} (ç¼–å· {num}): æ–‡æœ¬{len(info['text_fields'])}é¡¹, å›¾ç‰‡{len(info['image_fields'])}é¡¹"
        for num, info in templates.items()
    )
    image_section = "æ— " if not images else "\n".join(images)
    prompt = f"""
è¯·å°†ä»¥ä¸‹è®²ç¨¿æ‹†åˆ†æˆè‹¥å¹²å¼  PPTï¼Œæ¯å¼ å¹»ç¯ç‰‡é€‰æ‹©ä¸€ä¸ªæ¨¡æ¿ï¼Œå¹¶è¾“å‡º JSON æ•°ç»„ï¼Œæ¯ä¸ªå…ƒç´ åŒ…å«ï¼š
- template_page_num: æ¨¡æ¿ç¼–å·
- page_type: æ¨¡æ¿åç§°
- texts: æŒ‰æ¨¡æ¿æ–‡æœ¬å­—æ®µé¡ºåºç»™å‡ºçš„å†…å®¹æ•°ç»„
- images: æŒ‰æ¨¡æ¿å›¾ç‰‡å­—æ®µé¡ºåºç»™å‡ºçš„å†…å®¹æ•°ç»„

æ¨¡æ¿ä¿¡æ¯ï¼š
{template_desc}

å¯ç”¨å›¾ç‰‡è·¯å¾„ï¼š
{image_section}

è¾“å‡ºæ ¼å¼ç¤ºä¾‹ï¼š
[
  {{
    "template_page_num": 4,
    "page_type": "ç›®å½•é¡µ",
    "texts": ["ç›®å½•æ ‡é¢˜", "æ¡ç›®1", "æ¡ç›®1è¯´æ˜", "..."],
    "images": [""]
  }},
  ...
]

ç”Ÿæˆå†…å®¹æ—¶åŠ¡å¿…è®°ä½å¹¶é‡å¤ä½¿ç”¨è®²ç¨¿ä¸­çš„ä¸»è®²äººå§“åã€è¯¾ç¨‹/è®²åº§/é¡¹ç›®åç§°ç­‰å…³é”®ä¸“æœ‰åè¯ï¼Œç¡®ä¿åœ¨æ‰€æœ‰å¹»ç¯ç‰‡ä¸­éœ€è¦å¡«å†™ä¸“æœ‰åè¯çš„ä½ç½®ä¿æŒä¸€è‡´ï¼Œä¸è¦éšæ„æ”¹å†™æˆ–å¦é€ æ–°åç§°ã€‚æ‰€æœ‰æ ‡è®°ä¸ºâ€œrequiredâ€çš„å­—æ®µéƒ½å¿…é¡»æä¾›æ–‡å­—ï¼Œä¸”ä¸å¾—è¶…è¿‡å¯¹åº”çš„ max_chars é™åˆ¶ã€‚
è®²ç¨¿å…¨æ–‡ï¼š
{doc_text}
"""
    response = llm.generate([{"role": "user", "content": prompt}], temperature=0.3)
    try:
        plan = _ensure_json_array(response)
        return plan
    except Exception:
        raise ValueError("æ¨¡å‹è¾“å‡ºæ— æ³•è§£æä¸º JSON æ•°ç»„ï¼Œè¯·æ£€æŸ¥æç¤ºæˆ–é‡è¯•ã€‚")


def _ensure_json_object(text: str) -> Dict:
    text = text.strip()
    start = text.find("{")
    end = text.rfind("}")
    if start == -1 or end == -1:
        raise ValueError("æ¨¡å‹è¾“å‡ºä¸­æœªæ‰¾åˆ° JSON å¯¹è±¡")
    return json.loads(text[start : end + 1])


def _ensure_json_array(text: str) -> List[Dict]:
    text = text.strip()
    start = text.find("[")
    end = text.rfind("]")
    if start == -1 or end == -1:
        raise ValueError("æ¨¡å‹è¾“å‡ºä¸­æœªæ‰¾åˆ° JSON æ•°ç»„")
    return json.loads(text[start : end + 1])


def choose_llm(enable: bool, provider: str, model: Optional[str]) -> Optional[BaseLLM]:
    if not enable:
        return None
    provider = (provider or "").lower()
    if provider == "deepseek":
        return DeepSeekLLM(model=model or "deepseek-chat")
    if provider == "local":
        return LocalLLM(model=model)
    raise ValueError(f"æš‚ä¸æ”¯æŒçš„å¤§æ¨¡å‹æä¾›å•†ï¼š{provider}")


def _apply_metadata_overrides(content: Dict, template_info: Dict, metadata: Dict):
    if not metadata:
        return
    for field in template_info["text_fields"]:
        path = list(field["path"])
        key = "/".join(path)
        value = None
        if metadata.get("lecturer") and "ä¸»è®²" in key:
            value = metadata["lecturer"]
        elif metadata.get("college") and "å­¦é™¢" in key:
            value = metadata["college"]
        elif metadata.get("course") and any(
            token in key for token in ("è¯¾ç¨‹", "è¯¾ç¨‹åç§°", "é¡¹ç›®")
        ):
            value = metadata["course"]
        if value is not None:
            _assign_in_schema(content, path, value)


def _fill_with_template(
    template_num: int,
    template_info: Dict,
    block: Dict,
    llm: Optional[BaseLLM],
    metadata: Dict,
) -> Dict:
    content = llm_fill_slide(
        llm, template_info, block.get("text", ""), block.get("images", [])
    )
    _apply_metadata_overrides(content, template_info, metadata)
    return {
        "page_type": template_info["page_type"],
        "template_page_num": template_num,
        "content": content,
    }


def _strip_values(node):
    if isinstance(node, dict):
        if "type" in node and "value" in node:
            return node.get("value", "")
        return {k: _strip_values(v) for k, v in node.items()}
    if isinstance(node, list):
        return [_strip_values(item) for item in node]
    return node


def _empty_content(template_info: Dict) -> Dict:
    result = _clone_schema(template_info["schema"])
    for field in template_info["text_fields"]:
        _assign_in_schema(result, list(field["path"]), "")
    for field in template_info["image_fields"]:
        _assign_in_schema(result, list(field["path"]), "")
    return result


def _prepend_cover_page(pages: List[Dict], templates: Dict[int, Dict]):
    cover_template = templates.get(1)
    if not cover_template:
        return
    if pages and pages[0].get("template_page_num") == 1:
        return
    pages.insert(
        0,
        {
            "page_type": cover_template["page_type"],
            "template_page_num": 1,
            "content": _empty_content(cover_template),
        },
    )


def _fill_by_markers(
    blocks: List[Dict],
    templates: Dict[int, Dict],
    llm: Optional[BaseLLM],
    metadata: Dict,
) -> List[Dict]:
    pages: List[Dict] = []
    for block in blocks:
        template_num = block.get("template_hint")
        if template_num is None:
            continue
        if template_num not in templates:
            raise ValueError(
                f"æ¨¡æ¿ {template_num} æœªåœ¨ template.json ä¸­å®šä¹‰æˆ–ä¸åœ¨ template.txt ä¸­å…è®¸ã€‚"
            )
        pages.append(
            _fill_with_template(
                template_num, templates[template_num], block, llm, metadata
            )
        )
    return pages


def _plan_without_markers(
    blocks: List[Dict], templates: Dict[int, Dict], llm: BaseLLM, metadata: Dict
) -> List[Dict]:
    if not llm:
        raise ValueError("è®²ç¨¿æœªæŒ‡å®š PPT æ ‡è®°ä¸”æœªå¯ç”¨ LLMï¼Œæ— æ³•è‡ªåŠ¨åˆ†é…æ¨¡æ¿ã€‚")
    doc_text = "\n\n".join(
        block.get("text", "") for block in blocks if block.get("text")
    )
    all_images = [path for block in blocks for path in block.get("images", [])]
    plan = llm_plan_slides(llm, doc_text, templates, all_images)
    pages: List[Dict] = []
    for item in plan:
        template_num = item.get("template_page_num")
        if template_num not in templates:
            raise ValueError(f"æ¨¡å‹è¿”å›çš„æ¨¡æ¿ç¼–å· {template_num} ä¸åœ¨å…è®¸åˆ—è¡¨ä¸­ã€‚")
        template_info = templates[template_num]
        content = _clone_schema(template_info["schema"])
        texts = item.get("texts", [])
        images = item.get("images", [])
        for idx, field in enumerate(template_info["text_fields"]):
            value = _lookup_field_value(
                field, texts, texts if isinstance(texts, list) else None, idx
            )
            _assign_in_schema(content, list(field["path"]), value)
        for idx, field in enumerate(template_info["image_fields"]):
            value = _lookup_field_value(
                field, images, images if isinstance(images, list) else None, idx
            )
            _assign_in_schema(content, list(field["path"]), value)
        _apply_metadata_overrides(content, template_info, metadata)
        pages.append(
            {
                "page_type": template_info["page_type"],
                "template_page_num": template_num,
                "content": content,
            }
        )
    return pages


def generate_config_data(
    docx_path: str,
    template_json: str,
    template_list: str,
    use_llm: bool,
    llm_provider: str,
    llm_model: Optional[str],
    metadata_overrides: Optional[Dict[str, str]],
    run_dir: Path,
) -> Dict:
    """æ ¸å¿ƒé€»è¾‘ï¼šç”Ÿæˆ JSON å†…å®¹ï¼Œä¾› GUI/CLI å¤ç”¨ã€‚"""
    metadata_overrides = metadata_overrides or {}
    image_dir = run_dir / "images"
    blocks, has_marker, metadata = parse_docx_blocks(docx_path, image_dir)
    for key in ("course", "college", "lecturer"):
        if metadata_overrides.get(key):
            metadata[key] = metadata_overrides[key]

    templates = load_template_defs(template_json, template_list)
    llm = choose_llm(use_llm, llm_provider, llm_model)

    if has_marker:
        pages = _fill_by_markers(blocks, templates, llm, metadata)
    else:
        pages = _plan_without_markers(blocks, templates, llm, metadata)

    if not pages:
        raise ValueError("æœªç”Ÿæˆä»»ä½•å¹»ç¯ç‰‡å†…å®¹ï¼Œè¯·æ£€æŸ¥è®²ç¨¿æˆ–æ¨¡æ¿ã€‚")

    _prepend_cover_page(pages, templates)

    stripped_pages = []
    for page in pages:
        stripped_pages.append(
            {
                "page_type": page.get("page_type"),
                "template_page_num": page.get("template_page_num"),
                "content": _strip_values(page.get("content", {})),
            }
        )

    return {"ppt_pages": stripped_pages}


def process_docx(
    docx_path: str,
    template_json: str,
    template_list: str,
    output_path: Optional[str],
    use_llm: bool,
    llm_provider: str,
    llm_model: Optional[str],
    override_course: Optional[str],
    override_college: Optional[str],
    override_lecturer: Optional[str],
    run_dir: Optional[str],
    config_name: str,
):
    """CLI åŒ…è£…ï¼šå¤„ç†å‚æ•°ã€ä¿è¯ run ç›®å½•å­˜åœ¨ï¼Œå¹¶é¢å¤–å¤åˆ¶æ–‡ä»¶åˆ° outputã€‚"""
    metadata_overrides = {
        "course": override_course,
        "college": override_college,
        "lecturer": override_lecturer,
    }

    base_dir = Path(run_dir) if run_dir else _create_run_dir()
    base_dir.mkdir(parents=True, exist_ok=True)
    config_path = base_dir / config_name

    config = generate_config_data(
        docx_path,
        template_json,
        template_list,
        use_llm,
        llm_provider,
        llm_model,
        metadata_overrides,
        base_dir,
    )

    config_path.write_text(json.dumps(config, ensure_ascii=False, indent=2), encoding="utf-8")

    if output_path:
        explicit = Path(output_path)
        explicit.parent.mkdir(parents=True, exist_ok=True)
        shutil.copyfile(config_path, explicit)
        print(f"ğŸ“„ å¦å­˜ä¸ºï¼š{explicit}")

    print(f"âœ… å·²ç”Ÿæˆ JSONï¼š{config_path}")
    print(f"ğŸ“ èµ„æºè¾“å‡ºç›®å½•ï¼š{base_dir}")


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="æ ¹æ® DOCX è®²ç¨¿ç”Ÿæˆ PPT é…ç½® JSONã€‚")
    parser.add_argument("--docx", required=True, help="è®²ç¨¿ DOCX è·¯å¾„")
    parser.add_argument(
        "--template-json", default="template/template.json", help="æ¨¡æ¿å®šä¹‰ JSON æ–‡ä»¶"
    )
    parser.add_argument(
        "--template-list", default="template/template.txt", help="å¯ç”¨æ¨¡æ¿ç¼–å·åˆ—è¡¨ txt"
    )
    parser.add_argument(
        "--output",
        default=None,
        help="å¦‚éœ€é¢å¤–å¤åˆ¶ä¸€ä»½ JSONï¼Œè¯·æä¾›å®Œæ•´è·¯å¾„ï¼›è‹¥çœç•¥åˆ™ä»…åœ¨ temp/run-*/ ä¸­ç”Ÿæˆ",
    )
    parser.add_argument("--use-llm", action="store_true", help="å¯ç”¨å¤§æ¨¡å‹å¡«å……/æ’ç‰ˆ")
    parser.add_argument("--llm-provider", default="deepseek", help="å¤§æ¨¡å‹æä¾›å•†")
    parser.add_argument("--llm-model", default="deepseek-chat", help="å¤§æ¨¡å‹åç§°")
    parser.add_argument("--course-name", default=None, help="æ‰‹åŠ¨æŒ‡å®šè¯¾ç¨‹/é¡¹ç›®åç§°")
    parser.add_argument("--college-name", default=None, help="æ‰‹åŠ¨æŒ‡å®šå­¦é™¢/å•ä½")
    parser.add_argument("--lecturer-name", default=None, help="æ‰‹åŠ¨æŒ‡å®šä¸»è®²æ•™å¸ˆå§“å")
    parser.add_argument(
        "--run-dir",
        default=None,
        help="æŒ‡å®šè¾“å‡ºç›®å½•ï¼ˆé»˜è®¤åœ¨ temp ä¸‹è‡ªåŠ¨åˆ›å»º run-æ—¶é—´æˆ³-éšæœºå€¼ æ–‡ä»¶å¤¹ï¼‰",
    )
    parser.add_argument(
        "--config-name",
        default="config.json",
        help="è¾“å‡ºç›®å½•ä¸­ç”Ÿæˆçš„é…ç½®æ–‡ä»¶åç§°",
    )
    return parser


def main():
    args = build_arg_parser().parse_args()
    process_docx(
        docx_path=args.docx,
        template_json=args.template_json,
        template_list=args.template_list,
        output_path=args.output,
        use_llm=args.use_llm,
        llm_provider=args.llm_provider,
        llm_model=args.llm_model,
        override_course=args.course_name,
        override_college=args.college_name,
        override_lecturer=args.lecturer_name,
        run_dir=args.run_dir,
        config_name=args.config_name,
    )


def _strip_values(node):
    if isinstance(node, dict):
        if "type" in node and "value" in node:
            return node.get("value", "")
        return {k: _strip_values(v) for k, v in node.items()}
    if isinstance(node, list):
        return [_strip_values(item) for item in node]
    return node


if __name__ == "__main__":
    main()
